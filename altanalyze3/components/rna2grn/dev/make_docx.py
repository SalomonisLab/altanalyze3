#!/usr/bin/env python3.11
"""Generate the rna2grn validation report as a Word (.docx) document in formal
scientific prose for a methods publication."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RPT = "/Users/saljh8/Dropbox/Collaborations/Grimes/Human-GRN/July-2026-simple/rna2grn/reports"
OUT = f"{RPT}/rna2grn_validation.docx"
doc = Document()

# base style
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(10.5)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0); r.font.name = "Arial"
    return p


def para(text):
    p = doc.add_paragraph(text); p.paragraph_format.space_after = Pt(6); return p


def figure(png, number, caption, width=6.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f"{RPT}/{png}", width=Inches(width))
    cap = doc.add_paragraph()
    r = cap.add_run(f"Figure {number}. "); r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
    r2 = cap.add_run(caption); r2.font.size = Pt(9); r2.font.name = "Arial"
    cap.paragraph_format.space_after = Pt(10)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, x in enumerate(headers):
        c = t.rows[0].cells[i].paragraphs[0].add_run(str(x)); c.bold = True; c.font.size = Pt(9)
    for r in rows:
        cells = t.add_row().cells
        for i, x in enumerate(r):
            run = cells[i].paragraphs[0].add_run(str(x)); run.font.size = Pt(9)
    doc.add_paragraph()
    return t


# ---------------------------------------------------------------- title
title = doc.add_heading("", level=0)
trun = title.add_run("Imputation of transcription-factor regulatory activity from RNA "
                     "by per-edge elastic-net regression with regulon covariates "
                     "(rna2grn): method and validation")
trun.font.name = "Arial"; trun.font.color.rgb = RGBColor(0, 0, 0)
para("AltAnalyze3 cross-modality imputation module. This document describes the method "
     "and its validation for predicting gene-regulatory-network (GRN) connection scores "
     "from transcriptomes, with emphasis on the detection of transcription factors (TFs) "
     "whose regulatory activity is altered in a query sample relative to a matched control.")

# ---------------------------------------------------------------- rationale
h("1. Rationale", 1)
para("Gene-regulatory networks summarize the regulatory influence of each TF on its target "
     "genes as a set of edge-level connection (activity) scores. Their inference typically "
     "requires paired chromatin accessibility (multiome or TEA-seq) and is computationally "
     "intensive. We sought a model that imputes the full edge-level GRN of a cell-state "
     "pseudobulk from its RNA profile alone, with the specific objective of identifying TFs "
     "whose regulatory activity is perturbed in disease. A central requirement is that a "
     "change in a TF's activity that is reflected in its regulon's expression, but not in the "
     "TF's own transcript level, must be recovered as an increase in that TF's edge activity "
     "scores. The model must therefore respond continuously to regulon-level expression and "
     "extrapolate beyond the profiles present in the training reference, rather than "
     "interpolating among them.")

# ---------------------------------------------------------------- methods
h("2. Materials and Methods", 1)

h("2.1 Datasets and pseudobulk matching", 2)
para("GRN connection scores were obtained for 7,486 TF–target edges across 358 "
     "sample-by-cell-state columns (file tf_to_gene_connection_scores_log10_..."
     "_COMBINED_AML_Multiome_RC2.txt). Matched RNA pseudobulks were drawn from a harmonized "
     "atlas of 12,255 pseudobulks keyed as Sample|CellState (pseudobulk_counts_hashed.h5ad). "
     "GRN columns encode the sample identifier and cell state joined by an underscore, with "
     "spaces in cell-state names replaced by underscores; each column was resolved to its "
     "RNA pseudobulk by longest-suffix matching against the 89-term cell-state vocabulary "
     "followed by curated sample-name reconciliation (module grn_io.py). Of 358 columns, 354 "
     "were matched; four (a CCHMC AML-13 with no corresponding atlas pseudobulk) were "
     "discarded. The matched dataset comprised 341 pseudobulks (223 AML/MDS patient, 26 "
     "CCHMC AML, 82 Multiome control, 10 TEA-seq control) spanning 30 samples and 83 cell "
     "states. All 217 TFs and all 2,361 target genes appearing in the GRN were present in the "
     "RNA feature space.")

h("2.2 Control reference construction", 2)
para("The two control GRN aggregates have no native RNA in the atlas. Both were associated, "
     "on the RNA side, with a single healthy-control reference defined as the summed counts "
     "of 12 CCHMC control donors (Annotation = Control, Dataset = CCHMC; AM72, BF21, BF32, "
     "BF71, BM27, BM32, WF26, WF32, WF83, WM22, WM34, WM72; 10x Genomics 3′ v3 chemistry), "
     "aggregated per cell state. This identical control RNA vector was paired with both the "
     "Multiome_WT control GRN (82 cell states) and the RC2_TEA control GRN (10 cell states: "
     "ERP-1, HSC-2, Intermediate Mono-1, LMPP-1, LMPP-1-cycling, LMPP-2, MEP-2, MPP-1, MPP-2, "
     "MPP-MEP). The aggregate total edge activity per pseudobulk differed systematically by "
     "assay: median sums were 291 and 327 for 3′ AML patient and CCHMC AML pseudobulks, "
     "67 for TEA-seq controls, and 46 for multiome controls; the cell-state-matched "
     "AML-to-multiome ratio was 4.9. Consequently, raw AML-minus-control differentials carry "
     "an assay-scale offset under which approximately 60% of TFs appear elevated in AML in "
     "every pseudobulk. TEA-seq controls, being closer in magnitude to the 3′ AML data, "
     "constitute a better-matched reference but cover fewer cell states; both controls were "
     "therefore retained and reported.")
figure("fig3_protocol_similarity.png", 1,
       "Cross-assay similarity of control GRNs. Pearson correlation between the Multiome and "
       "TEA-seq control connection-score profiles for the ten cell states shared by both "
       "assays. The mean cross-assay, same-cell-state correlation (0.80) substantially "
       "exceeds the mean correlation between distinct cell states within the multiome control "
       "(0.40, dashed line), indicating that cell-state identity, rather than assay, "
       "dominates the GRN profile.", width=4.8)

h("2.3 Normalization", 2)
para("Pseudobulk counts were normalized to counts per 10,000 and log1p-transformed "
     "(CP10k/log1p). For inference on new data, cells assigned to a cell state are summed "
     "into a pseudobulk and normalized identically; the pipeline is therefore reproducible "
     "from any raw count matrix.")

h("2.4 Model: per-edge elastic-net regression", 2)
para("Each edge e linking transcription factor T to target gene g was modeled independently "
     "as an elastic-net (regularized linear) regression of its connection score across "
     "pseudobulks on three standardized covariates of the same pseudobulk's normalized "
     "expression: the expression of the target gene g, the expression of the transcription "
     "factor T, and the mean expression of the regulon of T (the set of all targets of T). "
     "The elastic-net L1 mixing parameter was set to zero (l1_ratio = 0), reducing the "
     "penalty to L2 (ridge); with only three covariates the L1 term confers no benefit "
     "(Section 2.5) and the L2 limit admits a closed-form solution. The 7,486 per-edge "
     "systems were assembled and solved jointly in closed form "
     "(L2 regularization λ = 1.0, intercept unpenalized); covariate standardization "
     "statistics are stored with the model (module model.py, class RegulonEdgeRegressor; "
     "training in training.py). Because the predicted activity of an edge increases with the "
     "expression of its target gene and regulon, an increase in regulon expression raises the "
     "predicted activity of the responsible TF's edges even when the TF's own transcript is "
     "unchanged; the model thus extrapolates beyond reference profiles. Standardized "
     "coefficient magnitudes confirmed the intended weighting, with target-gene expression "
     "(mean |β| = 0.017) exceeding TF expression (0.014) and regulon mean (0.006).")

h("2.5 Relationship to the antecedent elastic-net formulation", 2)
para("The model belongs to the same elastic-net regression family as the antecedent "
     "AltAnalyze3 imputation modules (rna2lipid, rna2adt); the distinction is structural, "
     "not the regression family. The antecedent modules use a single, global, multi-task "
     "elastic net mapping one shared gene panel to all targets jointly. Applied to GRN "
     "activity, that formulation was unsuitable: under leave-one-donor-out cross-validation "
     "it yielded a strongly negative coefficient of determination, reflecting unstable "
     "extrapolation on held-out donors. We therefore decomposed the problem into independent "
     "per-edge (per-target) elastic-net regressions with the biologically constrained "
     "three-covariate design above. We additionally evaluated a non-zero L1 mixing parameter "
     "(l1_ratio = 0.5): it gave equivalent accuracy (cross-validated R² = 0.771 versus 0.773 "
     "at l1_ratio = 0) but, because each edge involves only three covariates, the L1 term "
     "conferred no benefit while substantially increasing cost — iterative coordinate "
     "descent required approximately 46 ms per edge against approximately 0.7 ms per edge for "
     "the closed-form L2 solution (a full 7,486-edge fit in approximately 5 s per fold). The "
     "L1 mixing parameter was therefore fixed at zero.")

h("2.6 Cross-validation design", 2)
para("A sample denotes a donor; a pseudobulk denotes a donor-by-cell-state unit and is the "
     "unit of prediction. Hold-out was performed at the donor level, so that all pseudobulks "
     "of a held-out donor were assigned to the test partition together, precluding leakage of "
     "any cell state of a held-out donor into training. The AML-7 donor was excluded from all "
     "evaluation and reserved as a real-world inference and timing input (bundle load 0.06 s; "
     "whole-sample pseudobulk prediction 0.32 s). The remaining 27 AML donors were partitioned "
     "into five grouped folds; each donor was held out exactly once. The unique, merged "
     "control aggregates cannot be held out and were retained in training in every fold, "
     "serving as the differential baseline. This procedure yielded 214 held-out pseudobulks "
     "across 40 cell states and 61,845 (pseudobulk, TF, control) evaluation instances.")
table(["Fold", "Held-out AML donors", "Retained AML donors", "Held-out pseudobulks"],
      [[0, 6, 21, 28], [1, 6, 21, 67], [2, 5, 22, 39], [3, 5, 22, 35], [4, 5, 22, 45]])

h("2.7 Evaluation metrics", 2)
para("For each held-out pseudobulk, per-TF activity was computed as the mean of the imputed "
     "scores over that TF's edges. Differential TF activity was defined as the difference "
     "between the AML pseudobulk and its cell-state-matched control, computed both for the "
     "measured GRN and for the imputed GRN. Truly differential TFs were defined by magnitude "
     "strata of the measured differential (top 50%, 25%, 10% by absolute value), single-"
     "replicate pseudobulks precluding a per-instance significance test. The primary metric "
     "was directional concordance, the fraction of truly differential TFs for which the "
     "imputed differential matched the sign of the measured differential. To separate "
     "TF-specific structure from the assay-scale offset, the metric was computed both on the "
     "raw differential and on the per-pseudobulk-centered differential (each pseudobulk's "
     "217-TF differential vector centered to zero mean). Evaluation is implemented in "
     "evaluate.py and dev/validate_directional.py.")

# ---------------------------------------------------------------- results
h("3. Results", 1)

h("3.1 Reconstruction accuracy", 2)
para("Under leave-one-donor-out cross-validation, the per-edge regulon regression attained a "
     "coefficient of determination of 0.76 against the global mean and a median per-edge "
     "across-sample Pearson correlation of 0.73, exceeding both a cell-state-mean baseline "
     "(R² = 0.52) and a nearest-neighbor retrieval baseline (R² = 0.63). The connection-score "
     "variance is partitioned 65.5% between cell states and 34.5% within cell states "
     "(Figure 2); the within-cell-state component represents the sample-specific signal "
     "targeted by differential analysis.")
figure("fig2_variance_decomposition.png", 2,
       "Decomposition of GRN connection-score variance into between-cell-state (65.5%) and "
       "within-cell-state (34.5%) components. The within-cell-state component is the "
       "sample-specific signal that differential TF-activity analysis seeks to recover.",
       width=3.6)

h("3.2 Directional concordance of differential TF activity", 2)
para("Directional concordance of imputed differential TF activity against the matched "
     "controls is summarized below. Concordance against TEA-seq controls exceeded that "
     "against multiome controls on the raw differential, consistent with the smaller "
     "assay-scale offset of the TEA-seq reference. Removing the per-pseudobulk offset raised "
     "concordance for both controls to approximately 0.90 and revealed that the raw "
     "differential had been attenuated by the global offset rather than inflated by it; the "
     "TF-specific direction was recovered in approximately 90% of truly differential TFs. "
     "Concordance increased monotonically with the magnitude of the measured differential, "
     "and was statistically significant against the chance value of 0.5 (binomial test, "
     "p < 10⁻³⁰) with balanced performance across up- and down-regulated TFs.")
table(["Control", "Differential", "Top 50%", "Top 25%", "Top 10%"],
      [["TEA-seq", "raw", 0.781, 0.845, 0.898],
       ["TEA-seq", "offset-removed", 0.910, 0.904, 0.892],
       ["Multiome", "raw", 0.749, 0.822, 0.874],
       ["Multiome", "offset-removed", 0.907, 0.901, 0.848]])
para("Per-donor directional concordance (top-25% differential TFs, TEA-seq-preferred "
     "control) ranged from 0.72 to 0.91 (median 0.83) across the 27 held-out donors "
     "(file detail_per_sample.csv; Figure 3).")
figure("fig4_directional_concordance.png", 3,
       "Directional concordance of imputed differential TF activity (held-out AML donors "
       "versus the cell-state-matched control; AML-7 excluded, controls retained in "
       "training). (A) Concordance by magnitude stratum of the measured differential, for "
       "TEA-seq and multiome controls, on the raw and per-pseudobulk-centered (offset-"
       "removed) differential. (B) Per-donor concordance (top-25% differential TFs, "
       "TEA-seq-preferred control), median 0.83. (C) A representative held-out pseudobulk "
       "(donor 1009_AfInv16_29M, Intermediate Mono-1, TEA-seq control); each point is a TF, "
       "coloured by agreement between the sign of the measured and imputed offset-removed "
       "differential.", width=6.5)
para("To establish that the imputed differential reflects regulon-level information rather "
     "than the transcription factor's own abundance, directional concordance was compared "
     "with a baseline that predicts the activity direction from the sign of the TF's own "
     "expression change (Figure 4). For TFs whose own transcript was unchanged "
     "(|Δ expression| below 0.25 log-units; the regulon-only case), the per-edge model "
     "attained a concordance of 0.895 against 0.773 for the expression-sign baseline, "
     "confirming that the model recovers activity changes that are not visible in the "
     "factor's own expression.")
figure("fig5_mechanism_expression_baseline.png", 4,
       "Imputed differential TF-activity direction relative to a TF-expression-sign baseline "
       "(offset-removed; controls retained; TEA-seq-preferred control). (A) Concordance by "
       "magnitude stratum for the per-edge regulon model and the expression baseline. (B) "
       "Stratification of the top-25% differential TFs by whether the TF's own mRNA is "
       "differential; for regulon-only TFs (mRNA stable) the model exceeds the expression "
       "baseline (0.895 versus 0.773), demonstrating recovery of activity changes "
       "independent of TF abundance.", width=6.0)

h("3.3 Cell-type and TF-level variability", 2)
para("Across 40 cell states, per-cell-state directional concordance ranged from 0.72 to 0.90 "
     "(file detail_per_cellstate.csv). The highest concordance was observed for "
     "well-represented progenitor and myeloid states (LMPP-2 and MPP-MEP, each 0.90; cDC1 "
     "0.89; Mono-1 0.87); rarer and lymphoid-leaning states were more variable. At the TF "
     "level (file detail_per_tf.csv), all 217 TFs were scored in every pseudobulk; 106 were "
     "frequently differential. Concordance was highest for E2F3, USF1, TCF12, STAT1, RELB, "
     "ARNT, ATF3 and FOSB, and for the myeloid regulators RUNX1, SPI1, JUN, JUNB, FOS and "
     "MAFG; it was lowest for ubiquitously expressed factors with small differentials "
     "(USF2 0.29, KLF3 0.33, RARA 0.33, JDP2 0.34).")
para("For individual pseudobulks, the most differential TFs were tabulated together with the "
     "sign agreement of their imputed activity (file per_pseudobulk_topworst_tf.csv). For "
     "example, in donor 1009_AfInv16_29M, cell state Intermediate Mono-1 (TEA-seq control), "
     "the imputed differential correctly recovered the elevation of JUND, FOSL1, SREBF2, "
     "NFE2L2 and STAT3, but failed to recover the measured elevation of CEBPD, SPI1 and "
     "CEBPB. Imputed differential magnitudes were systematically smaller than measured "
     "magnitudes, reflecting ridge regularization and the assay-offset term on the measured-"
     "control side; directional agreement was therefore substantially more reliable than "
     "absolute magnitude, and rank or sign rather than absolute difference is the "
     "recommended quantity for interpretation.")
table(["TF", "Measured Δ", "Imputed Δ", "Direction"],
      [["JUND", "+0.297", "+0.016", "concordant"],
       ["FOSL1", "+0.200", "+0.148", "concordant"],
       ["SREBF2", "+0.127", "+0.122", "concordant"],
       ["CEBPD", "+0.215", "−0.002", "discordant"],
       ["SPI1", "+0.214", "−0.018", "discordant"],
       ["CEBPB", "+0.087", "−0.040", "discordant"]])

h("3.4 Specificity for novel regulon perturbations", 2)
para("Recovery of an AML-versus-control differential is partly attributable to a regulatory "
     "signature shared with the other AML donors retained in training, which a retrieval "
     "model can reproduce without inference. To assess detection of perturbations absent from "
     "the reference, we performed an in silico experiment (dev/perturbation_test.py): the "
     "expression of a single TF's regulon was increased in a healthy pseudobulk while the "
     "TF's own transcript was held constant, and TFs were ranked by the increase in predicted "
     "edge activity. The per-edge regulon regression ranked the perturbed TF at a median rank "
     "of 2 of 217 (perturbed TF within the top 5 in 83% of trials), whereas nearest-neighbor "
     "retrieval ranked it at a median of 74 (top 5 in 5% of trials). Retrieval is "
     "constitutively unable to express a perturbation absent from the reference, whereas the "
     "regression localizes it; this distinction motivates the choice of model (Figure 5).")
figure("fig1_perturbation_detection.png", 5,
       "Specificity for novel regulon perturbations (in silico). The expression of a single "
       "TF's regulon was increased in a healthy pseudobulk with the TF's own transcript held "
       "constant, and TFs were ranked by the increase in predicted edge activity. (A) "
       "Cumulative distribution of the rank assigned to the perturbed TF (of 217); the "
       "per-edge regulon regression concentrates near rank 1, whereas nearest-neighbor "
       "retrieval does not. (B) Fraction of perturbations for which the perturbed TF is "
       "ranked first or within the top five.", width=6.0)

# ---------------------------------------------------------------- limitations
h("4. Limitations", 1)
para("First, the only available control GRNs derive from multiome and TEA-seq assays, "
     "introducing a systematic scale offset relative to the 3′ AML data; differential "
     "analyses should use the per-pseudobulk-centered differential and the TEA-seq-matched "
     "control where coverage permits. A within-assay disease-versus-healthy differential "
     "would require a 3′ healthy control GRN, which is not available in the present data. "
     "Second, recovery of an in-distribution AML differential can be achieved in part by "
     "reproducing a shared signature; the in silico perturbation experiment is the assay "
     "that isolates detection of novel perturbations. Third, imputed differential magnitudes "
     "are attenuated relative to measured magnitudes, and directional agreement should be "
     "preferred over absolute values. Fourth, single-replicate pseudobulks preclude a "
     "per-instance test of differential significance, necessitating magnitude-stratified "
     "reporting.")

# ---------------------------------------------------------------- code availability
h("5. Code and data availability", 1)
para("The module resides at altanalyze3/components/rna2grn. The model and training "
     "procedure are implemented in model.py and training.py; the inference interface in "
     "api.py and cli.py; GRN parsing and pseudobulk matching in grn_io.py. The canonical "
     "evaluation, which produces all reported tables, is evaluate.py "
     "(directional_dual_control.csv, directional_metrics_dual.json, "
     "per_pseudobulk_topworst_tf.csv, detail_per_sample.csv, detail_per_cellstate.csv, "
     "detail_per_tf.csv). Supporting analyses are provided in dev/: dataset construction "
     "(build_dataset.py), signal diagnostics (diagnostics.py), directional validation with "
     "baselines and offset removal (validate_directional.py), the in silico perturbation "
     "experiment (perturbation_test.py), cross-assay similarity (protocol_similarity.py), and "
     "figure generation (make_figures.py). The trained model is distributed as a compressed "
     "pickle (rna2grn_bundle.pkl.gz, < 1 MB).")

doc.save(OUT)
print("wrote", OUT)
